from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, List
import os
import json
import qrcode
from datetime import datetime
from num2words import num2words

# Librerías para generación de PDF
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Librería para generación de Word (.docx)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI(
    title="API Integral de Gestión Documental y Búsqueda para Guinea Ecuatorial",
    version="4.0",
    description="API completa con búsqueda de servicios, plantilla de modelos, generador de PDF y Word con Código QR y formato administrativo local."
)

# ==========================================
# DIRECTORIOS Y CONFIGURACIÓN INICIAL
# ==========================================
DIR_TEMP = "temp_files"
DIR_MODELOS = "modelos_guardados"
DIR_LOGOS = "logos_guardados"

os.makedirs(DIR_TEMP, exist_ok=True)
os.makedirs(DIR_MODELOS, exist_ok=True)
os.makedirs(DIR_LOGOS, exist_ok=True)

RUTA_REGISTRO_MODELOS = os.path.join(DIR_MODELOS, "registro.json")

def cargar_registro():
    if os.path.exists(RUTA_REGISTRO_MODELOS):
        with open(RUTA_REGISTRO_MODELOS, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def guardar_registro(datos):
    with open(RUTA_REGISTRO_MODELOS, "w", encoding="utf-8") as f:
        json.dump(datos, f, ensure_ascii=False, indent=2)

def generar_qr(contenido: str, ruta_salida: str):
    """Genera una imagen QR de validación."""
    qr = qrcode.QRCode(version=1, box_size=4, border=1)
    qr.add_data(contenido)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(ruta_salida)

# ==========================================
# 1. BÚSQUEDA COMERCIAL Y DE INFORMACIÓN
# ==========================================
BASE_DATOS_MOCK = [
    {"id": 1, "tipo": "producto", "nombre": "Generador Eléctrico 5KW", "precio": 350000, "moneda": "XAF", "ubicacion": "Malabo"},
    {"id": 2, "tipo": "servicio", "nombre": "Mantenimiento de Climatización", "precio": 45000, "moneda": "XAF", "ubicacion": "Bata"},
    {"id": 3, "tipo": "servicio", "nombre": "Asesoría Legal y Tramitación de Visados", "precio": 150000, "moneda": "XAF", "ubicacion": "Malabo"},
]

@app.get("/buscar/comercial")
def buscar_productos_servicios(query: str, tipo: Optional[str] = None, max_precio: Optional[float] = None):
    """Busca productos y servicios filtrados de forma concreta."""
    resultados = BASE_DATOS_MOCK
    if query:
        resultados = [item for item in resultados if query.lower() in item["nombre"].lower()]
    if tipo:
        resultados = [item for item in resultados if item["tipo"].lower() == tipo.lower()]
    if max_precio:
        resultados = [item for item in resultados if item["precio"] <= max_precio]
    return {"total": len(resultados), "resultados": resultados}

class SolicitudInformacion(BaseModel):
    tema: str

@app.post("/buscar/informacion")
def buscar_informacion(solicitud: SolicitudInformacion):
    """Retorna los resultados categorizados por niveles de fiabilidad."""
    return {
        "tema_consultado": solicitud.tema,
        "jerarquia_fiabilidad": [
            {"nivel": 1, "tipo": "Gubernamental / Oficial (G.E. u Organismos Internacionales)", "peso": "100%"},
            {"nivel": 2, "tipo": "Académico / Técnico / Estadístico", "peso": "85%"},
            {"nivel": 3, "tipo": "Prensa Nacional e Internacional Reconocida", "peso": "70%"}
        ]
    }

# ==========================================
# 2. GESTIÓN DE MODELOS Y PLANTILLAS
# ==========================================
@app.post("/modelos/guardar")
async def guardar_modelo(
    nombre_modelo: str = Form(...),
    tipo_documento: str = Form(...),
    descripcion: str = Form(...),
    plantilla_texto: str = Form(...),
    logo: Optional[UploadFile] = File(None)
):
    """Guarda un modelo de referencia para reutilizarlo en la creación de documentos."""
    registro = cargar_registro()
    ruta_logo = None
    if logo:
        nombre_logo = f"{tipo_documento}_{logo.filename}"
        ruta_logo = os.path.join(DIR_LOGOS, nombre_logo)
        with open(ruta_logo, "wb") as f:
            f.write(await logo.read())
            
    registro[nombre_modelo] = {
        "tipo_documento": tipo_documento,
        "descripcion": descripcion,
        "plantilla_texto": plantilla_texto,
        "ruta_logo": ruta_logo,
        "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    guardar_registro(registro)
    return {"estado": "Éxito", "mensaje": f"Modelo '{nombre_modelo}' guardado correctamente."}

@app.get("/modelos/listar")
def listar_modelos():
    """Lista todos los modelos de documentos almacenados."""
    return cargar_registro()

# ==========================================
# 3. UTILIDAD DE MONTO EN LETRAS (FRANCOS CFA)
# ==========================================
@app.get("/utilidades/monto-letras")
def convertir_monto(monto: float):
    """Convierte importes numéricos a texto formal en Francos CFA."""
    letras = num2words(int(monto), lang='es').capitalize()
    return {"monto_numerico": monto, "monto_letras": f"{letras} Francos CFA"}

# ==========================================
# 4. GENERADOR MULTIFORMATO (PDF Y WORD)
# ==========================================
@app.post("/generar-documento")
async def generar_documento(
    formato: str = Form("pdf"),             # "pdf" o "word"
    tipo_documento: str = Form(...),        # "instancia", "solicitud", "factura", "cv"
    remitente: str = Form(...),
    destinatario: str = Form(...),
    cargo_destinatario: Optional[str] = Form(None),
    institucion_empresa: str = Form(...),
    ciudad: str = Form("Malabo"),
    asunto: str = Form(...),
    cuerpo_texto: str = Form(...),
    anexos: Optional[str] = Form(None),      # Documentos adjuntos
    nombre_modelo_usar: Optional[str] = Form(None),
    logo_o_foto: Optional[UploadFile] = File(None)
):
    """
    Genera un documento profesional en PDF o Word (.docx) adaptado al formato formal de Guinea Ecuatorial.
    """
    registro = cargar_registro()
    texto_final = cuerpo_texto
    ruta_logo_usar = None

    # Si se especifica un modelo previamente guardado, se fusiona
    if nombre_modelo_usar and nombre_modelo_usar in registro:
        mod = registro[nombre_modelo_usar]
        texto_final = f"{mod['plantilla_texto']}\n\nDETALLES ESPECÍFICOS:\n{cuerpo_texto}"
        if mod.get("ruta_logo"):
            ruta_logo_usar = mod["ruta_logo"]

    # Si el usuario sube un nuevo logo, sobrescribe el guardado
    if logo_o_foto:
        ruta_logo_usar = os.path.join(DIR_TEMP, logo_o_foto.filename)
        with open(ruta_logo_usar, "wb") as f:
            f.write(await logo_o_foto.read())

    fecha_actual = datetime.now().strftime("%d de %B de %Y")
    codigo_validacion = f"GE-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    
    # Generar QR de autenticidad
    ruta_qr = os.path.join(DIR_TEMP, f"qr_{codigo_validacion}.png")
    generar_qr(f"Validación de Documento Oficial G.E. | Código: {codigo_validacion} | Emisor: {remitente}", ruta_qr)

    # ------------------------------------------
    # GENERACIÓN EN FORMATO PDF
    # ------------------------------------------
    if formato.lower() == "pdf":
        nombre_pdf = f"{tipo_documento}_{codigo_validacion}.pdf"
        ruta_pdf = os.path.join(DIR_TEMP, nombre_pdf)
        
        doc = SimpleDocTemplate(ruta_pdf, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
        story = []
        styles = getSampleStyleSheet()

        # Logo opcional superior derecho
        if ruta_logo_usar and os.path.exists(ruta_logo_usar):
            img = RLImage(ruta_logo_usar, width=100, height=50)
            img.hAlign = 'RIGHT'
            story.append(img)
            story.append(Spacer(1, 10))

        estilo_titulo = ParagraphStyle('Titulo', parent=styles['Heading1'], fontSize=12, alignment=1, spaceAfter=15)
        estilo_cuerpo = ParagraphStyle('Cuerpo', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=10)
        estilo_negrita = ParagraphStyle('Negrita', parent=styles['Normal'], fontSize=10, leading=14, fontName='Helvetica-Bold')

        cargo_txt = cargo_destinatario.upper() if cargo_destinatario else 'AUTORIDAD COMPETENTE'
        
        story.append(Paragraph(f"<b>A LA ATENCIÓN DE:</b> {cargo_txt}", estilo_negrita))
        story.append(Paragraph(f"<b>{institucion_empresa.upper()}</b>", estilo_negrita))
        story.append(Paragraph(f"<b>{ciudad.upper()} (REPÚBLICA DE GUINEA ECUATORIAL)</b>", estilo_negrita))
        story.append(Spacer(1, 15))
        
        story.append(Paragraph(f"<b>ASUNTO:</b> {asunto.upper()}", estilo_titulo))
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("<b>EXCELENTÍSIMO / ILUSTRÍSIMO SEÑOR:</b>", estilo_negrita))
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Yo, <b>{remitente}</b>, comparezco y <b>EXPONGO:</b>", estilo_cuerpo))
        story.append(Paragraph(texto_final.replace('\n', '<br/>'), estilo_cuerpo))
        story.append(Spacer(1, 10))
        
        if anexos:
            story.append(Paragraph(f"<b>DOCUMENTACIÓN ADJUNTA (ANEXOS):</b><br/>{anexos.replace('\n', '<br/>')}", estilo_cuerpo))
            story.append(Spacer(1, 10))

        story.append(Paragraph("Por lo expuesto,", estilo_cuerpo))
        story.append(Paragraph("<b>SOLICITO</b> a V.E. / Ud. se sirva admitir el presente escrito y acceder a lo peticionado por ser de justicia.", estilo_cuerpo))
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"En {ciudad}, a {fecha_actual}.", estilo_cuerpo))
        story.append(Spacer(1, 30))
        
        # Bloque final con firma y Código QR
        qr_img = RLImage(ruta_qr, width=60, height=60)
        tabla_firma = Table(
            [[Paragraph(f"Firma / Sello:<br/><br/>_______________________<br/><b>{remitente}</b>", estilo_cuerpo), qr_img]],
            colWidths=[350, 100]
        )
        tabla_firma.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE')]))
        story.append(tabla_firma)

        doc.build(story)
        return FileResponse(path=ruta_pdf, filename=nombre_pdf, media_type='application/pdf')

    # ------------------------------------------
    # GENERACIÓN EN FORMATO WORD (.DOCX)
    # ------------------------------------------
    else:
        nombre_word = f"{tipo_documento}_{codigo_validacion}.docx"
        ruta_word = os.path.join(DIR_TEMP, nombre_word)
        
        doc_word = Document()
        
        if ruta_logo_usar and os.path.exists(ruta_logo_usar):
            p_logo = doc_word.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_logo.add_run().add_picture(ruta_logo_usar, width=Inches(1.5))

        p_dest = doc_word.add_paragraph()
        p_dest.add_run(f"A LA ATENCIÓN DE: {cargo_destinatario.upper() if cargo_destinatario else 'AUTORIDAD COMPETENTE'}\n").bold = True
        p_dest.add_run(f"{institucion_empresa.upper()}\n").bold = True
        p_dest.add_run(f"{ciudad.upper()} (REPÚBLICA DE GUINEA ECUATORIAL)").bold = True

        p_asunto = doc_word.add_paragraph()
        p_asunto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_asunto.add_run(f"\nASUNTO: {asunto.upper()}\n").bold = True

        doc_word.add_paragraph().add_run("EXCELENTÍSIMO / ILUSTRÍSIMO SEÑOR:").bold = True
        doc_word.add_paragraph(f"Yo, {remitente}, comparezco y EXPONGO:")
        doc_word.add_paragraph(texto_final)

        if anexos:
            p_anx = doc_word.add_paragraph()
            p_anx.add_run("DOCUMENTACIÓN ADJUNTA (ANEXOS):\n").bold = True
            p_anx.add_run(anexos)

        doc_word.add_paragraph("Por lo expuesto,")
        p_sol = doc_word.add_paragraph()
        p_sol.add_run("SOLICITO ").bold = True
        p_sol.add_run("a V.E. / Ud. se sirva admitir el presente escrito y acceder a lo peticionado por ser de justicia.")

        doc_word.add_paragraph(f"\nEn {ciudad}, a {fecha_actual}.")
        
        p_firma = doc_word.add_paragraph(f"\nFirma / Sello:\n\n_______________________\n{remitente}")
        
        doc_word.save(ruta_word)
        return FileResponse(path=ruta_word, filename=nombre_word, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
